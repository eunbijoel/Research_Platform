"""Research / meeting notes: summary → form fields → DOCX/HWPX export.

Adapted from Document_Analyser (writer + summarizer + DOCX/HWPX builders).
"""

from __future__ import annotations

import io
import re
from typing import Sequence

from research_memory.engine.llm import LLMConnectionError, generate_text, llm_available

MODE_RESEARCH = "research"
MODE_MEETING = "meeting"
NOTE_MODES = (MODE_RESEARCH, MODE_MEETING)

MODE_LABELS = {
    MODE_RESEARCH: "연구노트",
    MODE_MEETING: "회의록",
}

AUDIO_EXTENSIONS = {".mp3", ".wav", ".m4a", ".webm", ".ogg", ".flac", ".aac"}

RESEARCH_NOTE_SECTIONS = """
아래 형식을 그대로 지켜 작성하세요. 각 제목은 반드시 줄의 시작에 쓰고, 마크다운(#, **)은 쓰지 마세요.
원문에 없는 내용은 만들지 말고, 없으면 「확인 필요」라고 쓰세요.

주제 제안:

연구 또는 작업 목적

사용한 자료 및 파일

내용:

연구 결과:

확인된 문제점

향후 작업 계획
"""

MEETING_NOTE_SECTIONS = """
아래 형식을 그대로 지켜 작성하세요. 각 제목은 반드시 줄의 시작에 쓰고, 마크다운(#, **)은 쓰지 마세요.
녹음 트랜스크립트·회의 자료를 근거로만 작성하고, 없으면 「확인 필요」라고 쓰세요.
Action Item은 가능하면 「담당 / 기한 / 할 일」 한 줄씩 적어 주세요.

회의 제목:

일시:

참석자:

안건:

논의내용:

결정사항:

Action Item:
"""

_SECTION_SPECS: list[tuple[str, str]] = [
    ("topic", r"(?m)^\s*주제\s*제안\s*:?\s*"),
    ("purpose", r"(?m)^\s*연구\s*또는\s*작업\s*목적\s*:?\s*"),
    ("materials", r"(?m)^\s*사용한\s*자료(?:\s*및\s*파일)?\s*:?\s*"),
    ("content", r"(?m)^\s*(?:주요\s*)?내용\s*:?\s*"),
    ("results", r"(?m)^\s*연구\s*결과\s*:?\s*"),
    ("issues", r"(?m)^\s*확인된\s*문제점\s*:?\s*"),
    ("plan", r"(?m)^\s*향후\s*작업\s*계획\s*:?\s*"),
]

_MEETING_SECTION_SPECS: list[tuple[str, str]] = [
    ("topic", r"(?m)^\s*회의\s*제목\s*:?\s*"),
    ("datetime", r"(?m)^\s*일\s*시\s*:?\s*"),
    ("attendees", r"(?m)^\s*참석자\s*:?\s*"),
    ("agenda", r"(?m)^\s*안\s*건\s*:?\s*"),
    ("discussion", r"(?m)^\s*논의\s*내용\s*:?\s*"),
    ("decisions", r"(?m)^\s*결정\s*사항\s*:?\s*"),
    ("actions", r"(?m)^\s*Action\s*Item\s*:?\s*"),
]


def normalize_note_mode(value: str | None) -> str:
    raw = (value or "").strip().lower()
    if raw in {"meeting", "회의록", "minutes", "meeting_minutes"}:
        return MODE_MEETING
    return MODE_RESEARCH


def _parse_sections(
    text: str,
    specs: list[tuple[str, str]],
    out_keys: tuple[str, ...],
    *,
    fallback_key: str,
) -> dict[str, str]:
    raw = (text or "").replace("\r\n", "\n").strip()
    out = {k: "" for k in out_keys}
    if not raw:
        return out

    hits: list[tuple[int, str, int]] = []
    for key, pat in specs:
        m = re.search(pat, raw)
        if m:
            hits.append((m.start(), key, m.end()))
    if not hits:
        out[fallback_key] = raw
        return out

    hits.sort(key=lambda x: x[0])
    for i, (_start, key, end) in enumerate(hits):
        if key not in out:
            continue
        stop = hits[i + 1][0] if i + 1 < len(hits) else len(raw)
        body = raw[end:stop].strip()
        body = re.sub(r"^\s*[-·*]\s*", "", body, count=1)
        out[key] = body.strip()
    return out


def parse_research_note_fields(text: str) -> dict[str, str]:
    """통합 요약에서 주제/내용/연구결과만 추출."""
    return _parse_sections(
        text,
        _SECTION_SPECS,
        ("topic", "content", "results"),
        fallback_key="content",
    )


def parse_meeting_note_fields(text: str) -> dict[str, str]:
    """회의록 요약에서 표 필드 추출."""
    return _parse_sections(
        text,
        _MEETING_SECTION_SPECS,
        ("topic", "datetime", "attendees", "agenda", "discussion", "decisions", "actions"),
        fallback_key="discussion",
    )


def suggest_title(filenames: list[str], summary: str, *, mode: str = MODE_RESEARCH) -> str:
    mode = normalize_note_mode(mode)
    prefix = MODE_LABELS[mode]
    if len(filenames) == 1:
        base = re.sub(r"\.[^.]+$", "", filenames[0])
        return f"{prefix} — {base}"
    if filenames:
        return f"{prefix} — {filenames[0]} 외 {len(filenames) - 1}건"
    m = re.search(r"(?:목적|주제|회의\s*제목)[:：]?\s*(.{4,40})", summary)
    if m:
        return f"{prefix} — {m.group(1).strip()}"
    return prefix


def generate_research_note_summary(
    source_text: str,
    *,
    filenames: list[str] | None = None,
    mode: str = MODE_RESEARCH,
) -> str:
    """LLM으로 연구노트/회의록용 통합 요약 생성. 실패 시 원문 앞부분을 반환."""
    text = (source_text or "").strip()
    if not text:
        return ""
    mode = normalize_note_mode(mode)
    names = filenames or []
    file_line = f"참고 파일: {', '.join(names)}\n\n" if names else ""
    if mode == MODE_MEETING:
        prompt = (
            "다음 자료(회의 녹음 트랜스크립트·메모·첨부)를 바탕으로 회의록 초안을 작성하세요.\n\n"
            f"{MEETING_NOTE_SECTIONS}\n\n"
            f"{file_line}"
            f"자료:\n{text[:12000]}"
        )
        fallback = _fallback_meeting_summary
    else:
        prompt = (
            "다음 자료를 바탕으로 연구노트용 통합 요약을 작성하세요.\n\n"
            f"{RESEARCH_NOTE_SECTIONS}\n\n"
            f"{file_line}"
            f"자료:\n{text[:8000]}"
        )
        fallback = _fallback_summary
    if not llm_available():
        return fallback(text, names)
    try:
        out = generate_text(prompt).strip()
        return out or fallback(text, names)
    except LLMConnectionError:
        return fallback(text, names)


def _fallback_summary(text: str, filenames: list[str]) -> str:
    title = suggest_title(filenames, text, mode=MODE_RESEARCH)
    snippet = text.strip().replace("\n", " ")
    if len(snippet) > 600:
        snippet = snippet[:597] + "…"
    return (
        f"주제 제안:\n{title}\n\n"
        f"연구 또는 작업 목적:\n확인 필요\n\n"
        f"사용한 자료 및 파일:\n{', '.join(filenames) if filenames else '확인 필요'}\n\n"
        f"내용:\n{snippet}\n\n"
        f"연구 결과:\n확인 필요\n\n"
        f"확인된 문제점:\n확인 필요\n\n"
        f"향후 작업 계획:\n확인 필요"
    )


def _fallback_meeting_summary(text: str, filenames: list[str]) -> str:
    title = suggest_title(filenames, text, mode=MODE_MEETING)
    snippet = text.strip().replace("\n", " ")
    if len(snippet) > 800:
        snippet = snippet[:797] + "…"
    return (
        f"회의 제목:\n{title}\n\n"
        f"일시:\n확인 필요\n\n"
        f"참석자:\n확인 필요\n\n"
        f"안건:\n확인 필요\n\n"
        f"논의내용:\n{snippet}\n\n"
        f"결정사항:\n확인 필요\n\n"
        f"Action Item:\n확인 필요"
    )


def is_audio_filename(name: str) -> bool:
    from pathlib import Path

    return Path(name or "").suffix.lower() in AUDIO_EXTENSIONS


_whisper_model = None


def stt_available() -> bool:
    try:
        import faster_whisper  # noqa: F401

        return True
    except ImportError:
        pass
    try:
        import whisper  # noqa: F401

        return True
    except ImportError:
        return False


def transcribe_audio_bytes(data: bytes, filename: str = "audio.wav") -> tuple[str, str]:
    """Best-effort local STT. Returns (transcript, error).

    Prefers ``faster_whisper``, then ``openai-whisper``.
    """
    if not data:
        return "", "빈 오디오 파일입니다."
    if not stt_available():
        return (
            "",
            "자동 받아쓰기를 쓸 수 없습니다. "
            "프로젝트 venv에 `pip install faster-whisper` 후 앱을 재시작하거나, "
            "트랜스크립트를 직접 붙여넣어 주세요.",
        )
    try:
        import tempfile
        from pathlib import Path

        suffix = Path(filename).suffix.lower() or ".wav"
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(data)
            path = tmp.name
        try:
            err = ""
            try:
                text = _transcribe_with_faster_whisper(path)
            except Exception as exc:  # noqa: BLE001
                text, err = "", f"faster-whisper 변환 실패: {exc}"
            if text:
                return text, ""
            try:
                text = _transcribe_with_whisper(path)
            except Exception as exc:  # noqa: BLE001
                text, err = "", f"whisper 변환 실패: {exc}"
            if text:
                return text, ""
            return "", err or "음성을 텍스트로 변환하지 못했습니다. 파일 형식·길이를 확인해 주세요."
        finally:
            Path(path).unlink(missing_ok=True)
    except Exception as exc:  # noqa: BLE001
        return "", f"음성 변환 실패: {exc}"


def _transcribe_with_faster_whisper(path: str) -> str:
    global _whisper_model
    from faster_whisper import WhisperModel  # type: ignore

    if _whisper_model is None:
        _whisper_model = WhisperModel("base", device="cpu", compute_type="int8")
    segments, _info = _whisper_model.transcribe(path, language="ko")
    return "\n".join(
        seg.text.strip() for seg in segments if seg.text and seg.text.strip()
    ).strip()


def _transcribe_with_whisper(path: str) -> str:
    import whisper  # type: ignore

    model = whisper.load_model("base")
    result = model.transcribe(path, language="ko")
    return str((result or {}).get("text") or "").strip()


def note_rows(
    *,
    topic: str,
    owner: str,
    date_s: str,
    author: str,
    content: str,
    results: str,
    etc: str,
) -> list[tuple[str, str]]:
    return [
        ("주 제", topic or ""),
        ("책 임 자", owner or ""),
        ("일 시", date_s or ""),
        ("작 성 자", author or ""),
        ("내 용", content or ""),
        ("연구결과", results or ""),
        ("기타내용", etc or ""),
    ]


def meeting_rows(
    *,
    topic: str,
    date_s: str,
    attendees: str,
    agenda: str,
    discussion: str,
    decisions: str,
    actions: str,
) -> list[tuple[str, str]]:
    return [
        ("회의 제목", topic or ""),
        ("일 시", date_s or ""),
        ("참석자", attendees or ""),
        ("안 건", agenda or ""),
        ("논의내용", discussion or ""),
        ("결정사항", decisions or ""),
        ("Action Item", actions or ""),
    ]


def note_as_markdown(
    *,
    topic: str,
    owner: str,
    date_s: str,
    author: str,
    content: str,
    results: str,
    etc: str,
    project_id: str = "",
) -> str:
    """Plain markdown twin of the research-note table (good for Memory ingest)."""
    lines = [
        f"# 연구노트 — {topic or '제목 없음'}",
        "",
    ]
    if project_id:
        lines.append(f"Project: {project_id}")
        lines.append("")
    for label, val in note_rows(
        topic=topic,
        owner=owner,
        date_s=date_s,
        author=author,
        content=content,
        results=results,
        etc=etc,
    ):
        lines.append(f"## {label.strip()}")
        lines.append(val.strip() or "확인 필요")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def meeting_as_markdown(
    *,
    topic: str,
    date_s: str,
    attendees: str,
    agenda: str,
    discussion: str,
    decisions: str,
    actions: str,
    project_id: str = "",
    recording_name: str = "",
) -> str:
    """Plain markdown twin of the meeting-minutes table."""
    lines = [
        f"# 회의록 — {topic or '제목 없음'}",
        "",
    ]
    if project_id:
        lines.append(f"Project: {project_id}")
        lines.append("")
    if recording_name:
        lines.append(f"Recording: {recording_name}")
        lines.append("")
    for label, val in meeting_rows(
        topic=topic,
        date_s=date_s,
        attendees=attendees,
        agenda=agenda,
        discussion=discussion,
        decisions=decisions,
        actions=actions,
    ):
        lines.append(f"## {label.strip()}")
        lines.append(val.strip() or "확인 필요")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def _normalize_note_rows(rows: Sequence[tuple[str, str]]) -> list[tuple[str, str]]:
    out = [(str(label or ""), str(val or "")) for label, val in rows]
    return out or [("내 용", " ")]


def _strip_md_noise(text: str) -> str:
    out_lines: list[str] = []
    for ln in (text or "").splitlines():
        s = ln.rstrip()
        s = re.sub(r"^#{1,6}\s*", "", s)
        s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
        s = re.sub(r"__(.+?)__", r"\1", s)
        s = re.sub(r"`([^`]+)`", r"\1", s)
        if re.fullmatch(r"-{3,}|_{3,}|\*{3,}", s.strip()):
            continue
        out_lines.append(s)
    return "\n".join(out_lines)


def build_research_note_docx(rows: Sequence[tuple[str, str]], *, title: str = "연구노트") -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    note_rows_list = _normalize_note_rows(rows)
    d = Document()
    h = d.add_paragraph(title)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if h.runs:
        h.runs[0].bold = True
        h.runs[0].font.size = Pt(16)

    table = d.add_table(rows=len(note_rows_list), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Cm(3.2)
    table.columns[1].width = Cm(13.0)

    for i, (label, val) in enumerate(note_rows_list):
        c0, c1 = table.rows[i].cells
        c0.text = label
        c1.text = _strip_md_noise(val)
        try:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F0F0F0")
            shading.set(qn("w:val"), "clear")
            c0._tc.get_or_add_tcPr().append(shading)
        except Exception:
            pass
        for p in c0.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def build_research_note_hwpx(rows: Sequence[tuple[str, str]], *, title: str = "연구노트") -> bytes:
    """Optional HWPX export when python-hwpx is installed."""
    from hwpx import HwpxDocument  # type: ignore

    note_rows_list = _normalize_note_rows(rows)
    doc = HwpxDocument.new()
    paras = list(getattr(doc, "paragraphs", None) or [])
    if paras:
        try:
            paras[0].text = title
        except Exception:
            doc.add_paragraph(title)
    else:
        doc.add_paragraph(title)

    table = doc.add_table(rows=len(note_rows_list), cols=2, width=45_000)
    try:
        table.set_column_widths([2200, 7800])
    except Exception:
        pass
    for i, (label, val) in enumerate(note_rows_list):
        text = val if str(val).strip() else " "
        table.set_cell_text(i, 0, label if label else " ")
        table.set_cell_text(i, 1, _strip_md_noise(text), split_paragraphs=True)
        try:
            table.set_cell_shading(i, 0, "#F0F0F0")
        except Exception:
            pass
    raw = doc.to_bytes()
    if not raw or raw[:2] != b"PK":
        raise RuntimeError("연구노트 HWPX 생성 실패")
    return raw


def hwpx_available() -> bool:
    try:
        import hwpx  # noqa: F401

        return True
    except Exception:
        return False
