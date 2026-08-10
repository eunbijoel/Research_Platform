from __future__ import annotations

import json
import re
from datetime import datetime
from io import BytesIO
from typing import Any

from docx import Document

from research_memory.engine.llm import LLMConnectionError, generate_text, llm_available
from research_memory.engine.retrieval import retrieve
from research_memory.kb.repository import KnowledgeRepository
from research_memory.pipeline.chunking import refine_chunks
from research_memory.pipeline.extractors import extract_chunks_from_bytes
from research_memory.schema import (
    ROLE_PROJECT,
    ROLE_REFERENCE,
    Citation,
    normalize_document_role,
)

NOT_FOUND = "확인 필요"

RFP_FIELDS = [
    "project_name",
    "purpose",
    "organization",
    "duration",
    "budget",
    "mandatory_requirements",
    "tech_requirements",
    "kpi",
    "consortium_conditions",
    "evaluation_criteria",
    "submission_documents",
    "notes",
]

DRAFT_KEYS = [
    "necessity",
    "center_role",
    "work_details",
    "yearly_plan",
    "deliverables",
    "kpi_draft",
    "consortium_role",
    "expected_effects",
    "compliance_notes",
    "open_questions",
]

# Appended after narrative sections (template table — not LLM-filled amounts).
BUDGET_PLAN_KEY = "budget_plan"

DRAFT_LABELS = {
    "necessity": "참여 필요성",
    "center_role": "담당 역할",
    "work_details": "세부 수행내용",
    "yearly_plan": "연차별 수행계획",
    "deliverables": "예상 산출물",
    "kpi_draft": "KPI 초안",
    "consortium_role": "컨소시엄 내 역할",
    "expected_effects": "기대효과",
    "compliance_notes": "운영요령·참고규정 준수 포인트",
    "open_questions": "추가 확인이 필요한 사항",
    "budget_plan": "다음 단계 연구개발비 사용 계획",
}

# Section-aware draft: which evidence pool + keywords to re-rank (no new retrieval).
SECTION_SPECS: dict[str, dict[str, Any]] = {
    "necessity": {
        "use_research": True,
        "use_reference": False,
        "research_limit": 6,
        "reference_limit": 0,
        "rfp_fields": ["purpose", "project_name", "notes"],
        "keywords": [
            "배경", "필요", "목적", "정책", "산업", "문제", "동향", "도전", "기회",
            "기술주권", "온프레미스", "AIDC",
        ],
        "instruction": (
            "참여 필요성만 작성. 보고용 개조식(불릿 또는 짧은 명사형 문장)으로 3~6항목. "
            "합니다/습니다체 금지. 근거 문장 나열·복붙 금지."
        ),
    },
    "center_role": {
        "use_research": True,
        "use_reference": False,
        "research_limit": 8,
        "reference_limit": 0,
        "rfp_fields": ["mandatory_requirements", "tech_requirements", "consortium_conditions"],
        "keywords": [
            "역할", "담당", "센터", "총괄", "아키텍처", "연구노트", "제안서", "보고서",
            "역량", "모듈", "책임",
        ],
        "instruction": (
            "담당 역할만 작성. 개조식으로 역할 범위·경계를 3~6항목. "
            "합니다/습니다체 금지. '총괄'은 근거에 있을 때만, 기본은 주관·담당·역할 분담."
        ),
    },
    "work_details": {
        "use_research": True,
        "use_reference": False,
        "research_limit": 8,
        "reference_limit": 0,
        "rfp_fields": ["mandatory_requirements", "tech_requirements"],
        "keywords": [
            "수행", "구축", "개발", "설계", "구현", "검증", "기술", "시스템", "모듈",
            "아키텍처", "표준", "제안",
        ],
        "instruction": (
            "세부 수행내용만 작성. 1. 2. 3. 과업 단위 개조식. "
            "각 항목은 '~함/~개발/~구축' 명사형. 합니다/습니다체·출처 나열 금지."
        ),
    },
    "yearly_plan": {
        "use_research": True,
        "use_reference": False,
        "research_limit": 6,
        "reference_limit": 0,
        "rfp_fields": ["duration", "mandatory_requirements", "kpi"],
        "keywords": ["연차", "단계", "일정", "로드맵", "1차", "2차", "마일스톤", "기간"],
        "instruction": (
            "연차별 수행계획만 작성. 연차/단계별 개조식 목록. "
            "합니다/습니다체 금지. 근거 부족 시 '확인 필요'만 표기."
        ),
    },
    "deliverables": {
        "use_research": True,
        "use_reference": False,
        "research_limit": 6,
        "reference_limit": 0,
        "rfp_fields": ["submission_documents", "mandatory_requirements", "kpi"],
        "keywords": [
            "산출물", "보고서", "결과물", "패키지", "매뉴얼", "설계서", "프로토타입",
            "deliverable",
        ],
        "instruction": (
            "예상 산출물만 작성. '- 산출물명: 한 줄 설명' 개조식 목록. "
            "합니다/습니다체·출처 태그 금지."
        ),
    },
    "kpi_draft": {
        "use_research": True,
        "use_reference": False,
        "research_limit": 5,
        "reference_limit": 0,
        "rfp_fields": ["kpi", "evaluation_criteria", "tech_requirements"],
        "keywords": ["KPI", "지표", "성능", "목표", "평가", "달성", "정량"],
        "instruction": (
            "KPI 초안만 작성. '- 지표명: 목표(수치)' 개조식. "
            "합니다/습니다체 금지. 수치 근거 없으면 확인 필요."
        ),
    },
    "consortium_role": {
        "use_research": True,
        "use_reference": False,
        "research_limit": 5,
        "reference_limit": 0,
        "rfp_fields": ["consortium_conditions", "mandatory_requirements"],
        "keywords": ["컨소시엄", "주관", "참여기관", "협력", "분담", "인터페이스", "역할"],
        "instruction": (
            "컨소시엄 내 역할만 작성. 개조식으로 책임 경계·인터페이스 정리. "
            "합니다/습니다체 금지. '총괄'은 근거에 있을 때만."
        ),
    },
    "expected_effects": {
        "use_research": True,
        "use_reference": False,
        "research_limit": 5,
        "reference_limit": 0,
        "rfp_fields": ["purpose", "kpi", "notes"],
        "keywords": ["효과", "성과", "확산", "기대", "기여", "파급", "활용"],
        "instruction": (
            "기대효과만 작성. 기술·산업 효과를 개조식 불릿으로 구분. "
            "합니다/습니다체 금지."
        ),
    },
    "compliance_notes": {
        "use_research": False,
        "use_reference": True,
        "research_limit": 0,
        "reference_limit": 8,
        "rfp_fields": ["budget", "duration", "consortium_conditions"],
        "keywords": [
            "운영요령", "연구개발비", "계상", "제재", "참여제한", "성과관리", "현물",
            "인건비", "집행", "증빙",
        ],
        "instruction": (
            "운영요령·참고규정 준수 포인트만 작성. 계상·증빙·제재 등 사전 확인 사항을 "
            "개조식 불릿 3~6개. 합니다/습니다체 금지. 준수 완료 단정 금지."
        ),
    },
    "open_questions": {
        "use_research": True,
        "use_reference": True,
        "research_limit": 3,
        "reference_limit": 3,
        "rfp_fields": ["notes", "consortium_conditions", "budget", "kpi"],
        "keywords": ["확인", "미정", "협의", "분담", "예산", "범위"],
        "instruction": (
            "추가 확인 사항만 '- … 확인 필요' 개조식 체크리스트로 나열. "
            "합니다/습니다체·대괄호 태그 금지."
        ),
    },
}

_REFERENCE_QUERIES = [
    "산업기술혁신사업 공통 운영요령",
    "연구개발비 계상 산정 기준",
    "연구시설 장비 현물 계상",
    "성과관리 보고 의무",
    "제재 참여제한",
    "회의비 인건비 사용 제한",
]


def parse_rfp_bytes(data: bytes, filename: str) -> tuple[list[dict[str, str]], str]:
    """Extract citeable RFP chunks from an uploaded file."""
    _ftype, raw, err = extract_chunks_from_bytes(data, filename)
    if err and not raw:
        return [], err
    chunks = refine_chunks(raw)
    out = [
        {"file": filename, "location": c.location, "text": c.text}
        for c in chunks
        if c.text.strip()
    ]
    return out, "" if out else (err or "No text extracted from RFP")


def analyze_rfp(rfp_chunks: list[dict[str, str]]) -> dict[str, Any]:
    source = _chunks_to_text(rfp_chunks)
    if not source.strip():
        result = _empty_rfp()
        result["error"] = "RFP text is empty"
        return result

    if not llm_available():
        return _heuristic_rfp(source, rfp_chunks)

    prompt = f"""당신은 공공 R&D/용역 공고문을 분석하는 보조 도구입니다.
아래 [RFP 원문]에서만 근거를 찾아 JSON 객체 하나만 출력하세요.
설명 문장이나 마크다운 코드블록은 넣지 마세요.

각 항목:
- project_name, purpose, organization, duration, budget, consortium_conditions, notes: 문자열
- mandatory_requirements, tech_requirements, kpi, evaluation_criteria, submission_documents: 문자열 리스트
- 근거가 없으면 "{NOT_FOUND}" (리스트는 ["{NOT_FOUND}"])
- mandatory_requirements에는 실제 과업/수행내용만. 입찰자격은 notes 등으로.

[RFP 원문]
{source}
"""
    try:
        raw = generate_text(prompt)
        return _parse_rfp_json(raw)
    except LLMConnectionError as exc:
        result = _heuristic_rfp(source, rfp_chunks)
        result["error"] = str(exc)
        result["mode"] = "heuristic_fallback"
        return result


def gather_kb_evidence_split(
    rfp: dict[str, Any],
    *,
    repo: KnowledgeRepository | None = None,
    top_k_per_query: int = 4,
    project_id: str | None = None,
) -> dict[str, list[Citation]]:
    """
    Retrieve Memory evidence split by document_role.

    - research: project_document (optional project_id filter)
    - reference: reference_document (project filter ignored so 운영요령 survives)
    """
    repo = repo or KnowledgeRepository()
    research_queries = _research_queries(rfp)
    reference_queries = _reference_queries(rfp)

    research = _retrieve_filtered(
        research_queries,
        repo=repo,
        top_k_per_query=top_k_per_query,
        project_id=project_id,
        role=ROLE_PROJECT,
        prefer_regulation=False,
        limit=16,
    )
    reference = _retrieve_filtered(
        reference_queries,
        repo=repo,
        top_k_per_query=top_k_per_query,
        project_id=None,
        role=ROLE_REFERENCE,
        prefer_regulation=True,
        limit=12,
    )
    combined = _merge_citations(research + reference, limit=28)
    return {
        "research": research,
        "reference": reference,
        "combined": combined,
    }


def suggest_roles(
    rfp: dict[str, Any],
    evidence: list[Citation],
) -> list[dict[str, Any]]:
    evidence_text = _evidence_text(evidence, label="연구문서+참고규정")
    if not llm_available():
        return _heuristic_roles(rfp, evidence)

    prompt = f"""당신은 공공 R&D 컨소시엄 제안서 작성을 돕는 보조 도구입니다.
[RFP 분석]과 [Knowledge Base 근거]만 사용해 역할 후보를 최대 3개 제안하세요.
JSON 배열만 출력. 각 원소 키: role, reason, related_requirements, evidence, open_questions (문자열)

규칙:
- evidence는 KB 근거에 실제로 있는 내용만, 짧은 문서명을 함께 표시. 없으면 "{NOT_FOUND}"
- 근거 없는 추론은 추측임을 밝히되 [AI 제안] 같은 태그 문구는 쓰지 마세요.
- role 이름은 '총괄'을 기본으로 쓰지 마세요. RFP/근거에 명시된 경우만 '총괄'.
  기본은 주관·담당·참여·역할 분담 등 구체적 표현.
- 기술 역량은 연구 근거, 예산/의무/제재는 규정 근거를 구분해서 언급

[RFP 분석]
{json.dumps(rfp, ensure_ascii=False, indent=2)}

[Knowledge Base 근거]
{evidence_text}
"""
    try:
        raw = generate_text(prompt)
        return _parse_role_list(raw)
    except LLMConnectionError as exc:
        roles = _heuristic_roles(rfp, evidence)
        if roles:
            roles[0]["open_questions"] = str(exc)
        return roles


def select_section_evidence(
    section_key: str,
    rfp: dict[str, Any],
    research_evidence: list[Citation],
    reference_evidence: list[Citation],
) -> tuple[list[Citation], list[Citation]]:
    """
    Re-rank existing evidence pools for one draft section.
    Does not call retrieve — filters/scores the already-gathered citations.
    """
    spec = SECTION_SPECS.get(section_key) or {}
    terms = _section_terms(section_key, rfp, spec)
    research: list[Citation] = []
    reference: list[Citation] = []
    if spec.get("use_research", True):
        research = _rank_citations(
            research_evidence,
            terms=terms,
            limit=int(spec.get("research_limit") or 6),
        )
    if spec.get("use_reference", False):
        reference = _rank_citations(
            reference_evidence,
            terms=terms,
            limit=int(spec.get("reference_limit") or 6),
        )
    return research, reference


def generate_section_draft(
    section_key: str,
    rfp: dict[str, Any],
    selected_role: dict[str, Any],
    research_evidence: list[Citation],
    reference_evidence: list[Citation],
    *,
    review_notes: str | None = None,
) -> str:
    """Generate a single draft section string using section-scoped evidence."""
    if section_key not in DRAFT_KEYS:
        return NOT_FOUND
    label = DRAFT_LABELS.get(section_key, section_key)
    spec = SECTION_SPECS.get(section_key) or {}
    instruction = str(spec.get("instruction") or f"{label}만 작성하세요.")
    research_text = _evidence_text(research_evidence, label="연구 근거")
    reference_text = _evidence_text(reference_evidence, label="규정 근거")
    review_block = ""
    if (review_notes or "").strip():
        review_block = f"""
[초안 검토 피드백 — 이 섹션만 반영]
{(review_notes or "").strip()}
- 위 피드백을 반영해 이 섹션만 다시 작성.
- 보고용 개조식 유지(합니다/습니다 금지).
- 근거 없으면 새 사실 만들지 말고 '확인 필요'로 남김.
- 과도한 확정 표현(총괄·반드시·확정 등)은 완화.
"""

    if not llm_available():
        return _clean_draft_prose(
            _heuristic_section(
                section_key,
                rfp,
                selected_role,
                research_evidence,
                reference_evidence,
            )
        )

    prompt = f"""당신은 공공 R&D 제안서(보고서) 작성자입니다.
지금은 **하나의 섹션만** 작성합니다. JSON 객체만 출력하세요. 키는 정확히 "{section_key}" 하나 (값은 문자열).

섹션: {label} ({section_key})
지시: {instruction}
{review_block}
문체 (필수 — 보고용 개조식):
- '합니다/습니다/됩니다/입니다' 등 경어·서술체 금지.
- 명사형·개조식 종결(~함, ~임, ~필요, ~개발, ~구축, ~제시) 또는 '- 항목' 불릿 사용.
- 예: "온프레미스 AIDC 핵심기술 확보 필요" / "3계층 분산형 참조 아키텍처 설계·개발"
- 잘못된 예: "…확보가 필수적입니다." / "…역할을 수행할 것입니다."
- 근거 카드처럼 문장을 이어 붙이지 말고, 내용을 재구성한 개조식 항목으로 정리.
- [연구문서], [참고규정], [AI 제안], [확인 필요] 태그 금지.
- file=, location=, score=, section·part 표기 본문 금지.
- 출처가 꼭 필요하면 항목 끝에 (짧은문서명) 한 번만.
- 근거 없으면 새 사실 금지, '확인 필요'만 표기.
- '총괄'은 RFP/근거에 명시된 경우만. 기본은 주관·담당·역할 분담.
- 아래 근거는 내부 참고용. 근거 형식을 본문에 복사하지 말 것.

[RFP 분석]
{json.dumps(rfp, ensure_ascii=False, indent=2)}

[역할]
{json.dumps(selected_role, ensure_ascii=False, indent=2)}

[내부 참고 — 연구 근거]
{research_text}

[내부 참고 — 규정 근거]
{reference_text}
"""
    try:
        raw = generate_text(prompt)
        return _clean_draft_prose(_parse_section_text(raw, section_key))
    except LLMConnectionError:
        return _clean_draft_prose(
            _heuristic_section(
                section_key,
                rfp,
                selected_role,
                research_evidence,
                reference_evidence,
            )
        )


def generate_draft(
    rfp: dict[str, Any],
    selected_role: dict[str, Any],
    evidence: list[Citation] | None = None,
    *,
    research_evidence: list[Citation] | None = None,
    reference_evidence: list[Citation] | None = None,
) -> dict[str, Any]:
    """
    Section-aware draft: pick evidence per section, generate each section, then merge.
    Public signature unchanged for app.py / CLI.
    """
    research = list(research_evidence or [])
    reference = list(reference_evidence or [])
    if evidence and not research and not reference:
        research, reference = _split_by_role(evidence)

    draft: dict[str, Any] = {}
    used_research: list[Citation] = []
    used_reference: list[Citation] = []
    section_evidence: dict[str, Any] = {}
    errors: list[str] = []

    for key in DRAFT_KEYS:
        r_sub, f_sub = select_section_evidence(key, rfp, research, reference)
        section_evidence[key] = {
            "research": [c.to_dict() for c in r_sub],
            "reference": [c.to_dict() for c in f_sub],
        }
        used_research.extend(r_sub)
        used_reference.extend(f_sub)
        try:
            draft[key] = generate_section_draft(
                key, rfp, selected_role, r_sub, f_sub
            )
        except Exception as exc:  # noqa: BLE001 — keep other sections
            draft[key] = f"[확인 필요] 섹션 생성 실패: {exc}"
            errors.append(f"{key}: {exc}")

    used_research = _merge_citations(used_research, limit=40)
    used_reference = _merge_citations(used_reference, limit=24)
    combined = _merge_citations(used_research + used_reference, limit=48)
    draft["citations"] = [c.to_dict() for c in combined]
    draft["research_citations"] = [c.to_dict() for c in used_research]
    draft["reference_citations"] = [c.to_dict() for c in used_reference]
    draft["section_evidence"] = section_evidence
    draft["mode"] = "llm_section" if llm_available() else "heuristic_section"
    # Budget plan: empty stage-report skeleton (+ RFP total only). No auto-split.
    budget_md, budget_rows = build_budget_plan_section(rfp)
    draft[BUDGET_PLAN_KEY] = budget_md
    draft["budget_plan_table"] = budget_rows
    if errors:
        draft["error"] = "; ".join(errors)
    return draft


def build_budget_plan_section(rfp: dict[str, Any] | None = None) -> tuple[str, list[dict[str, str]]]:
    """Stage-report style budget skeleton. Amounts left blank except optional RFP total note."""
    rfp = rfp or {}
    total_note = _budget_total_note(rfp.get("budget"))
    cols = [
        "연구개발기관",
        "구분",
        "인건비(내부)",
        "인건비(외부)",
        "연구근접지원",
        "학생인건비(일반)",
        "학생인건비(특례)",
        "인건비소계",
        "시설·장비(일반)",
        "시설·장비(특례)",
        "연구재료비",
        "위탁연구개발비",
        "국제공동연구개발비",
        "연구활동비",
        "연구수당",
        "보안수당",
        "직접비소계",
        "간접비",
        "합계",
    ]
    row_specs: list[tuple[str, str]] = [
        ("주관연구개발기관", "현금"),
        ("주관연구개발기관", "현물"),
        ("주관연구개발기관", "소계"),
        ("주관연구개발기관", "미지급"),
        ("공동연구개발기관(A)", "현금"),
        ("공동연구개발기관(A)", "현물"),
        ("공동연구개발기관(A)", "소계"),
        ("공동연구개발기관(A)", "미지급"),
        ("공동연구개발기관(B)", "현금"),
        ("공동연구개발기관(B)", "현물"),
        ("공동연구개발기관(B)", "소계"),
        ("공동연구개발기관(B)", "미지급"),
        ("합계(주관+공동)", "현금"),
        ("합계(주관+공동)", "현물"),
        ("합계(주관+공동)", "합계(현금+현물)"),
        ("합계(주관+공동)", "미지급"),
        ("위탁연구개발기관(A)", "현금"),
        ("위탁연구개발기관(A)", "현물"),
        ("위탁연구개발기관(A)", "소계"),
        ("위탁연구개발기관(A)", "미지급"),
        ("위탁연구개발기관(B)", "현금"),
        ("위탁연구개발기관(B)", "현물"),
        ("위탁연구개발기관(B)", "소계"),
        ("위탁연구개발기관(B)", "미지급"),
        ("총계(주관+공동+위탁)", "현금"),
        ("총계(주관+공동+위탁)", "현물"),
        ("총계(주관+공동+위탁)", "총계(현금+현물)"),
        ("총계(주관+공동+위탁)", "미지급"),
    ]
    rows: list[dict[str, str]] = []
    for org, kind in row_specs:
        row = {c: "" for c in cols}
        row["연구개발기관"] = org
        row["구분"] = kind
        # Only annotate grand-total cash+in-kind with RFP total reference — never invent splits.
        if org.startswith("총계") and kind.startswith("총계") and total_note not in {
            "확인 필요",
            "",
        }:
            row["합계"] = total_note
        rows.append(row)

    header = "| " + " | ".join(cols) + " |"
    sep = "| " + " | ".join(["---"] * len(cols)) + " |"
    body_lines = []
    for row in rows:
        body_lines.append("| " + " | ".join(row.get(c, "") or "" for c in cols) + " |")

    md_parts = [
        "5) 다음 단계 연구개발비 사용 계획",
        "",
        "(단위 : 천원)",
        "",
        f"- RFP 예산(참고): {total_note}",
        "- 표는 단계보고서 양식 골격임. 기관·비목별 금액은 공란(자동 배분하지 않음).",
        "- 기관명·현금/현물/미지급·비목 금액은 확인 후 기입.",
        "",
        header,
        sep,
        *body_lines,
        "",
        "각주(개조식)",
        "- 1｣ 학생인건비 특례 미적용분",
        "- 2｣ 학생인건비 특례 적용분(영 제20조제4항제1호)",
        "- 3｣ 연구시설·장비비 특례 미적용분",
        "- 4｣ 연구시설·장비비 특례 적용분(영 제20조제4항제2호)",
        "- 5｣ 연구수당: 인건비(현물 포함, 연구근접지원인력 제외)+학생인건비 합계의 20% 이내",
        "- 6｣ 보안수당: 보안과제 참여연구자별 인건비의 3% 범위(해당 연구자만 지급)",
    ]
    return "\n".join(md_parts), rows


def _budget_total_note(raw: Any) -> str:
    text = str(raw or "").strip()
    if not text or text == NOT_FOUND:
        return "확인 필요"
    # Keep original text if no clear number; otherwise surface digits for the total cell.
    digits = re.sub(r"[^\d.]", "", text.replace(",", ""))
    if digits and re.search(r"\d", text):
        return f"{text} ※비목 배분은 확인 필요"
    return f"{text} ※비목 배분은 확인 필요"


REVIEW_CHECK_TYPES = (
    "rfp_gap",
    "unsupported_claim",
    "compliance_check",
    "inconsistency",
)

_LABEL_TO_KEY = {label: key for key, label in DRAFT_LABELS.items()}
_LABEL_TO_KEY.update(
    {
        "담당 역할": "center_role",
        "우리 센터의 담당 역할": "center_role",
        "수행내용": "work_details",
        "세부 수행내용": "work_details",
        "산출물": "deliverables",
        "예상 산출물": "deliverables",
        "운영요령": "compliance_notes",
        "운영요령·준수 포인트": "compliance_notes",
        "운영요령·참고규정 준수 포인트": "compliance_notes",
        "확인 필요": "open_questions",
        "추가 확인이 필요한 사항": "open_questions",
    }
)


def review_draft_quality(
    rfp: dict[str, Any],
    draft: dict[str, Any],
    *,
    research_evidence: list[Citation] | None = None,
    reference_evidence: list[Citation] | None = None,
    selected_role: dict[str, Any] | None = None,
) -> list[dict[str, Any]]:
    """Review Draft v1 against RFP + existing evidence. No new retrieval."""
    research = list(research_evidence or [])
    reference = list(reference_evidence or [])
    sections = {
        key: str(draft.get(key) or "").strip()
        for key in DRAFT_KEYS
        if str(draft.get(key) or "").strip()
    }
    if not sections:
        return [
            {
                "section": "전체",
                "check_type": "rfp_gap",
                "status": "확인 필요",
                "message": "검토할 Draft 섹션이 없습니다.",
                "evidence": [],
                "suggested_action": "Draft를 먼저 생성하세요.",
            }
        ]

    if not llm_available():
        return _heuristic_review(rfp, sections, research, reference)

    research_text = _evidence_text(research[:16], label="연구문서")
    reference_text = _evidence_text(reference[:16], label="참고규정")
    section_payload = {
        DRAFT_LABELS.get(k, k): v[:2500] for k, v in sections.items()
    }
    prompt = f"""당신은 공공 R&D 제안서 "센터 파트 초안"의 Quality Reviewer입니다.
새 사실을 만들지 말고, 아래 Draft를 RFP·연구문서·참고규정(운영요령) 기준으로만 검토하세요.
JSON 배열만 출력하세요.

각 원소 스키마:
{{
  "section": "섹션 한글명 (예: 참여 필요성, 담당 역할, 운영요령·참고규정 준수 포인트)",
  "check_type": "rfp_gap|unsupported_claim|compliance_check|inconsistency",
  "status": "문제없음|확인 필요",
  "message": "짧은 한국어 설명",
  "evidence": ["문서명 또는 RFP 필드/위치"],
  "suggested_action": "수정 방향 (확인 필요 표현 완화, RFP 확인 등)"
}}

검토 항목(이 4가지만):
1) rfp_gap — RFP 필수 요구·성능지표·산출물 등이 Draft에서 빠졌는지
2) unsupported_claim — Research/RFP 근거 없이 확정적으로 쓴 문장 ([확인 필요]로 가야 할 것)
3) compliance_check — 참고규정(특히 운영요령) 기준으로 사전 확인이 필요한지.
   법적/행정 준수 완료라고 단정하지 말고 "사전 검토/확인 필요" 톤으로.
4) inconsistency — 섹션 간 과한 중복, 숫자·역할·산출물 표현 충돌

규칙:
- status는 "문제없음" 또는 "확인 필요"만.
- 문제없는 주요 섹션도 가능하면 section당 1개 "문제없음" 항목을 포함.
- 근거가 없으면 evidence는 빈 배열, 추측으로 채우지 말 것.
- 참고규정은 Draft 작성 재료가 아니라 compliance 검토 기준으로 사용.

[RFP 분석]
{json.dumps(rfp, ensure_ascii=False, indent=2)}

[선택한 역할]
{json.dumps(selected_role or {}, ensure_ascii=False, indent=2)}

[Draft 섹션]
{json.dumps(section_payload, ensure_ascii=False, indent=2)}

[연구문서 근거]
{research_text}

[참고규정·운영요령 근거 — compliance 기준]
{reference_text}
"""
    try:
        raw = generate_text(prompt)
        findings = _parse_review_findings(raw)
        if findings:
            return findings
    except LLMConnectionError:
        pass
    return _heuristic_review(rfp, sections, research, reference)


def revise_draft_from_review(
    draft: dict[str, Any],
    findings: list[dict[str, Any]],
    rfp: dict[str, Any],
    selected_role: dict[str, Any],
    *,
    research_evidence: list[Citation] | None = None,
    reference_evidence: list[Citation] | None = None,
) -> dict[str, Any]:
    """Regenerate only sections marked 확인 필요; keep other Draft v1 text."""
    research = list(research_evidence or [])
    reference = list(reference_evidence or [])
    revised = dict(draft)
    notes_by_section: dict[str, list[str]] = {}
    for item in findings or []:
        if not isinstance(item, dict):
            continue
        if str(item.get("status") or "").strip() != "확인 필요":
            continue
        key = _section_label_to_key(str(item.get("section") or ""))
        if not key or key not in DRAFT_KEYS:
            continue
        msg = str(item.get("message") or "").strip()
        action = str(item.get("suggested_action") or "").strip()
        ctype = str(item.get("check_type") or "").strip()
        ev = item.get("evidence") or []
        ev_txt = ", ".join(str(x) for x in ev[:4]) if isinstance(ev, list) else str(ev)
        line = f"- ({ctype}) {msg}"
        if action:
            line += f" → {action}"
        if ev_txt:
            line += f" (근거 힌트: {ev_txt})"
        notes_by_section.setdefault(key, []).append(line)

    revised_keys: list[str] = []
    errors: list[str] = []
    for key, notes in notes_by_section.items():
        r_sub, f_sub = select_section_evidence(key, rfp, research, reference)
        review_notes = "\n".join(notes)
        try:
            revised[key] = generate_section_draft(
                key,
                rfp,
                selected_role,
                r_sub,
                f_sub,
                review_notes=review_notes,
            )
            revised_keys.append(key)
            # refresh section evidence snapshot for revised sections
            sec_ev = dict(revised.get("section_evidence") or {})
            sec_ev[key] = {
                "research": [c.to_dict() for c in r_sub],
                "reference": [c.to_dict() for c in f_sub],
            }
            revised["section_evidence"] = sec_ev
        except Exception as exc:  # noqa: BLE001
            revised[key] = f"[확인 필요] 섹션 개선 실패: {exc}"
            errors.append(f"{key}: {exc}")

    revised["mode"] = (
        "llm_section_revised" if llm_available() else "heuristic_section_revised"
    )
    revised["revised_sections"] = revised_keys
    # Keep / refresh budget skeleton (never LLM-split)
    budget_md, budget_rows = build_budget_plan_section(rfp)
    revised[BUDGET_PLAN_KEY] = budget_md
    revised["budget_plan_table"] = budget_rows
    if errors:
        revised["error"] = "; ".join(
            [str(revised.get("error") or "").strip(), *errors]
        ).strip("; ")
    return revised


def _section_label_to_key(label: str) -> str | None:
    text = (label or "").strip()
    if not text:
        return None
    if text in DRAFT_KEYS:
        return text
    if text in _LABEL_TO_KEY:
        return _LABEL_TO_KEY[text]
    # fuzzy contains
    for name, key in _LABEL_TO_KEY.items():
        if name in text or text in name:
            return key
    return None


def _parse_review_findings(raw: str) -> list[dict[str, Any]]:
    cleaned = _strip_fence(raw)
    try:
        parsed = json.loads(cleaned)
    except json.JSONDecodeError:
        # try array slice
        start = cleaned.find("[")
        end = cleaned.rfind("]")
        if start < 0 or end <= start:
            return []
        try:
            parsed = json.loads(cleaned[start : end + 1])
        except json.JSONDecodeError:
            return []
    if isinstance(parsed, dict):
        for key in ("findings", "results", "items", "reviews"):
            if isinstance(parsed.get(key), list):
                parsed = parsed[key]
                break
        else:
            return []
    if not isinstance(parsed, list):
        return []
    out: list[dict[str, Any]] = []
    for item in parsed:
        if not isinstance(item, dict):
            continue
        status = str(item.get("status") or "").strip()
        if status not in {"문제없음", "확인 필요"}:
            status = "확인 필요"
        ctype = str(item.get("check_type") or "").strip()
        if ctype not in REVIEW_CHECK_TYPES:
            ctype = "unsupported_claim"
        evidence = item.get("evidence") or []
        if not isinstance(evidence, list):
            evidence = [str(evidence)]
        out.append(
            {
                "section": str(item.get("section") or "").strip() or "전체",
                "check_type": ctype,
                "status": status,
                "message": str(item.get("message") or "").strip() or NOT_FOUND,
                "evidence": [str(x) for x in evidence[:6]],
                "suggested_action": str(item.get("suggested_action") or "").strip()
                or "확인 필요 표현으로 수정",
            }
        )
    return out


def _heuristic_review(
    rfp: dict[str, Any],
    sections: dict[str, str],
    research: list[Citation],
    reference: list[Citation],
) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    all_text = "\n".join(sections.values())
    # RFP gaps: keyword presence
    rfp_bits = []
    for field in ("mandatory_requirements", "kpi", "submission_documents"):
        val = rfp.get(field)
        if isinstance(val, list):
            rfp_bits.extend(str(x) for x in val[:5])
        elif val:
            rfp_bits.append(str(val))
    missing = []
    for bit in rfp_bits[:8]:
        token = bit.strip()[:20]
        if len(token) >= 4 and token not in all_text:
            missing.append(token)
    if missing:
        findings.append(
            {
                "section": DRAFT_LABELS.get("work_details", "세부 수행내용"),
                "check_type": "rfp_gap",
                "status": "확인 필요",
                "message": f"RFP 요구/지표 일부가 Draft에 명시적으로 연결되지 않았습니다: {', '.join(missing[:3])}",
                "evidence": ["RFP mandatory_requirements/kpi"],
                "suggested_action": "해당 요구를 수행내용에 매핑하거나 [확인 필요]로 표시",
            }
        )
    # Unsupported claims: assertive sentences without tags
    for key, text in sections.items():
        if key == "open_questions":
            continue
        if "[확인 필요]" in text or "[AI 제안]" in text or "확인 필요" in text:
            findings.append(
                {
                    "section": DRAFT_LABELS.get(key, key),
                    "check_type": "unsupported_claim",
                    "status": "문제없음",
                    "message": "불확실 표현(확인 필요)이 사용되어 있습니다.",
                    "evidence": [],
                    "suggested_action": "유지",
                }
            )
            continue
        if ("총괄" in text or "반드시" in text or "확정" in text) and (
            "(" not in text and "확인 필요" not in text
        ):
            findings.append(
                {
                    "section": DRAFT_LABELS.get(key, key),
                    "check_type": "unsupported_claim",
                    "status": "확인 필요",
                    "message": "확정·총괄 표현이 있으나 근거 표기가 약합니다.",
                    "evidence": [c.filename for c in research[:2]],
                    "suggested_action": "주관·담당으로 완화하거나 확인 필요로 표시",
                }
            )
    # Compliance
    if reference:
        comp = sections.get("compliance_notes") or ""
        if not comp or comp == NOT_FOUND or len(comp) < 40:
            findings.append(
                {
                    "section": DRAFT_LABELS.get("compliance_notes", "운영요령"),
                    "check_type": "compliance_check",
                    "status": "확인 필요",
                    "message": "참고규정(운영요령) 기준 사전 확인 항목이 Draft에 충분히 반영되지 않았습니다.",
                    "evidence": [c.filename for c in reference[:3]],
                    "suggested_action": "준수 완료 단정 없이 확인 필요 항목으로 보강",
                }
            )
        else:
            findings.append(
                {
                    "section": DRAFT_LABELS.get("compliance_notes", "운영요령"),
                    "check_type": "compliance_check",
                    "status": "확인 필요",
                    "message": "운영요령 관련 문구가 있으나 행정·예산 계상은 사전 검토가 필요합니다.",
                    "evidence": [c.filename for c in reference[:3]],
                    "suggested_action": "준수 완료로 단정하지 말고 확인 필요 톤 유지",
                }
            )
    else:
        findings.append(
            {
                "section": DRAFT_LABELS.get("compliance_notes", "운영요령"),
                "check_type": "compliance_check",
                "status": "확인 필요",
                "message": "참고규정 Evidence가 없어 운영요령 기준 검토를 완료할 수 없습니다.",
                "evidence": [],
                "suggested_action": "운영요령 문서를 Memory에 올린 뒤 재검토",
            }
        )
    # Inconsistency: repeated long phrases
    values = list(sections.values())
    if len(values) >= 2 and values[0][:80] and values[0][:80] in values[1]:
        findings.append(
            {
                "section": DRAFT_LABELS.get("center_role", "담당 역할"),
                "check_type": "inconsistency",
                "status": "확인 필요",
                "message": "섹션 간 동일 문구 반복이 감지되었습니다.",
                "evidence": [],
                "suggested_action": "역할/수행 경계를 구분해 중복 축소",
            }
        )
    if not findings:
        for key in list(sections.keys())[:3]:
            findings.append(
                {
                    "section": DRAFT_LABELS.get(key, key),
                    "check_type": "rfp_gap",
                    "status": "문제없음",
                    "message": "휴리스틱 검토에서 뚜렷한 문제는 보이지 않습니다.",
                    "evidence": [],
                    "suggested_action": "유지",
                }
            )
    return findings


def run_proposal_pipeline(
    rfp_data: bytes,
    rfp_filename: str,
    *,
    repo: KnowledgeRepository | None = None,
    project_id: str | None = None,
    selected_role_index: int = 0,
) -> dict[str, Any]:
    """End-to-end: parse → analyze → KB evidence → roles → draft."""
    repo = repo or KnowledgeRepository()
    chunks, err = parse_rfp_bytes(rfp_data, rfp_filename)
    if err and not chunks:
        return {"ok": False, "error": err}

    rfp = analyze_rfp(chunks)
    split = gather_kb_evidence_split(rfp, repo=repo, project_id=project_id)
    evidence = split["combined"]
    roles = suggest_roles(rfp, evidence)
    if not roles:
        roles = [
            {
                "role": NOT_FOUND,
                "reason": NOT_FOUND,
                "related_requirements": NOT_FOUND,
                "evidence": NOT_FOUND,
                "open_questions": "역할 후보 없음",
            }
        ]
    idx = max(0, min(selected_role_index, len(roles) - 1))
    selected = roles[idx]
    draft = generate_draft(
        rfp,
        selected,
        research_evidence=split["research"],
        reference_evidence=split["reference"],
    )
    md = build_markdown(
        rfp,
        selected,
        draft,
        research_evidence=split["research"],
        reference_evidence=split["reference"],
    )
    return {
        "ok": True,
        "rfp": rfp,
        "evidence": [c.to_dict() for c in evidence],
        "research_evidence": [c.to_dict() for c in split["research"]],
        "reference_evidence": [c.to_dict() for c in split["reference"]],
        "roles": roles,
        "selected_role": selected,
        "draft": draft,
        "markdown": md,
        "parse_error": err,
    }


def build_markdown(
    rfp: dict[str, Any],
    selected_role: dict[str, Any],
    draft: dict[str, Any],
    evidence: list[Citation] | list[dict[str, Any]] | None = None,
    *,
    research_evidence: list[Citation] | list[dict[str, Any]] | None = None,
    reference_evidence: list[Citation] | list[dict[str, Any]] | None = None,
) -> str:
    research = list(research_evidence or [])
    reference = list(reference_evidence or [])
    if evidence and not research and not reference:
        research = list(evidence)

    lines = [
        f"# 제안 초안 - {rfp.get('project_name', NOT_FOUND)}",
        "",
        f"_생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}_",
        "",
        "## RFP 요약",
        f"- 사업 목적: {rfp.get('purpose', NOT_FOUND)}",
        f"- 발주기관: {rfp.get('organization', NOT_FOUND)}",
        f"- 사업기간: {rfp.get('duration', NOT_FOUND)}",
        f"- 예산: {rfp.get('budget', NOT_FOUND)}",
        "",
        f"**역할:** {selected_role.get('role', NOT_FOUND)}",
        "",
    ]
    for key, label in DRAFT_LABELS.items():
        lines.append(f"### {label}")
        if key == BUDGET_PLAN_KEY:
            lines.append(str(draft.get(key) or "").strip() or NOT_FOUND)
        else:
            lines.append(clean_draft_prose(str(draft.get(key, NOT_FOUND))))
        lines.append("")

    if research or reference:
        lines.append("## 참고 근거")
        if research:
            lines.append("### 연구 자료")
            lines.extend(_format_evidence_lines(research, start=1))
        if reference:
            start = len(research) + 1
            lines.append("### 규정·운영요령")
            lines.extend(_format_evidence_lines(reference, start=start))

    return "\n".join(lines)


def export_docx_bytes(
    rfp: dict[str, Any],
    selected_role: dict[str, Any],
    draft: dict[str, Any],
) -> bytes:
    document = Document()
    document.add_heading(f"제안 초안 - {rfp.get('project_name', NOT_FOUND)}", level=0)
    document.add_paragraph(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    document.add_heading("RFP 요약", level=1)
    document.add_paragraph(f"사업 목적: {rfp.get('purpose', NOT_FOUND)}")
    document.add_paragraph(f"발주기관: {rfp.get('organization', NOT_FOUND)}")
    document.add_paragraph(f"사업기간: {rfp.get('duration', NOT_FOUND)}")
    document.add_paragraph(f"예산: {rfp.get('budget', NOT_FOUND)}")
    document.add_paragraph(f"역할: {selected_role.get('role', NOT_FOUND)}")
    for key, label in DRAFT_LABELS.items():
        document.add_heading(label, level=1)
        if key == BUDGET_PLAN_KEY:
            document.add_paragraph("(단위 : 천원) — 단계보고서 양식 골격. 비목 배분은 확인 필요.")
            rows = draft.get("budget_plan_table") or []
            if not rows:
                document.add_paragraph(str(draft.get(key) or NOT_FOUND)[:2000])
                continue
            cols = list(rows[0].keys())
            table = document.add_table(rows=1 + len(rows), cols=len(cols))
            table.style = "Table Grid"
            for j, col in enumerate(cols):
                table.rows[0].cells[j].text = str(col)
            for i, row in enumerate(rows, start=1):
                for j, col in enumerate(cols):
                    table.rows[i].cells[j].text = str(row.get(col) or "")
            document.add_paragraph(
                "각주: 1｣학생인건비 특례 미적용 2｣특례 적용 3｣시설·장비 특례 미적용 "
                "4｣특례 적용 5｣연구수당 20% 이내 6｣보안수당 인건비 3% 이내"
            )
            continue
        document.add_paragraph(clean_draft_prose(str(draft.get(key, NOT_FOUND))))
    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


# --- helpers ---

def clean_draft_prose(text: str) -> str:
    """Strip leftover citation tags / file= noise so sections read like proposal prose."""
    cleaned = (text or "").strip()
    if not cleaned:
        return cleaned
    # Bracket tags
    cleaned = re.sub(
        r"\[(?:연구문서|참고규정|AI 제안|확인 필요)\]\s*",
        "",
        cleaned,
    )
    # file=name | location=... ]  (including mismatched closing brackets)
    cleaned = re.sub(
        r"\bfile\s*=\s*([^\|\]]+?)(?:\s*\|\s*location\s*=\s*[^\]]*)?\]?",
        r"\1",
        cleaned,
        flags=re.IGNORECASE,
    )
    cleaned = re.sub(
        r"\s*\|\s*location\s*=\s*[^\]\|\n)]+",
        "",
        cleaned,
        flags=re.IGNORECASE,
    )
    cleaned = re.sub(
        r"\blocation\s*=\s*",
        "",
        cleaned,
        flags=re.IGNORECASE,
    )
    # Parenthetical dumps like (longname.hwpx | section 1 · part 3)
    cleaned = re.sub(
        r"\(([^()]+\.(?:hwpx|docx|pdf|xlsx|xls|txt|md))\s*\|\s*[^)]+\)",
        r"(\1)",
        cleaned,
        flags=re.IGNORECASE,
    )
    # Dangling "filename.ext]" leftovers
    cleaned = re.sub(
        r"(\S+\.(?:hwpx|docx|pdf|xlsx|xls|txt|md))\s*\]",
        r"\1",
        cleaned,
        flags=re.IGNORECASE,
    )
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _clean_draft_prose(text: str) -> str:
    return clean_draft_prose(text)


def _research_queries(rfp: dict[str, Any]) -> list[str]:
    queries: list[str] = []
    for key in ("project_name", "purpose"):
        val = rfp.get(key)
        if isinstance(val, str) and val and val != NOT_FOUND:
            queries.append(val)
    for key in ("mandatory_requirements", "tech_requirements", "kpi"):
        items = rfp.get(key) or []
        if isinstance(items, list):
            for item in items:
                if isinstance(item, str) and item and item != NOT_FOUND:
                    queries.append(item)
        elif isinstance(items, str) and items != NOT_FOUND:
            queries.append(items)
    if not queries:
        queries = ["연구 역량 제안서 산출물 센터 역할"]
    return queries[:12]


def _reference_queries(rfp: dict[str, Any]) -> list[str]:
    queries = list(_REFERENCE_QUERIES)
    for key in ("budget", "duration", "consortium_conditions"):
        val = rfp.get(key)
        if isinstance(val, str) and val and val != NOT_FOUND:
            queries.append(val)
    return queries[:10]


def _retrieve_filtered(
    queries: list[str],
    *,
    repo: KnowledgeRepository,
    top_k_per_query: int,
    project_id: str | None,
    role: str,
    prefer_regulation: bool,
    limit: int,
) -> list[Citation]:
    seen: set[tuple[str, str, str]] = set()
    citations: list[Citation] = []
    for q in queries:
        for c in retrieve(q, repo=repo, top_k=top_k_per_query):
            doc = repo.get_document(c.document_id) or {}
            doc_role = normalize_document_role(
                c.document_role or doc.get("document_role")
            )
            if doc_role != role:
                continue
            if project_id and (doc.get("project_id") or "") != project_id:
                continue
            # attach role/doc_type for downstream labeling
            c.document_role = doc_role
            key = (c.document_id, c.location, c.snippet[:80])
            if key in seen:
                continue
            seen.add(key)
            # slight boost for regulation docs when collecting reference evidence
            score = c.score
            if prefer_regulation and (doc.get("doc_type") or "").lower() == "regulation":
                score = float(score) + 0.01
            citations.append(
                Citation(
                    document_id=c.document_id,
                    filename=c.filename,
                    location=c.location,
                    snippet=c.snippet,
                    score=score,
                    document_role=doc_role,
                )
            )
    citations.sort(
        key=lambda x: (
            0
            if prefer_regulation
            and "운영요령" in (x.filename or "")
            else 1,
            -x.score,
        )
    )
    return citations[:limit]


def _merge_citations(items: list[Citation], *, limit: int) -> list[Citation]:
    seen: set[tuple[str, str, str]] = set()
    out: list[Citation] = []
    for c in items:
        key = (c.document_id, c.location, c.snippet[:80])
        if key in seen:
            continue
        seen.add(key)
        out.append(c)
        if len(out) >= limit:
            break
    return out


def _split_by_role(evidence: list[Citation]) -> tuple[list[Citation], list[Citation]]:
    research: list[Citation] = []
    reference: list[Citation] = []
    for c in evidence:
        if normalize_document_role(c.document_role) == ROLE_REFERENCE:
            reference.append(c)
        else:
            research.append(c)
    return research, reference


def _section_terms(section_key: str, rfp: dict[str, Any], spec: dict[str, Any]) -> list[str]:
    terms: list[str] = []
    for kw in spec.get("keywords") or []:
        if isinstance(kw, str) and kw.strip():
            terms.append(kw.strip().lower())
    for field in spec.get("rfp_fields") or []:
        val = rfp.get(field)
        if isinstance(val, list):
            for item in val:
                if isinstance(item, str) and item and item != NOT_FOUND:
                    terms.extend(_tokenize_terms(item))
        elif isinstance(val, str) and val and val != NOT_FOUND:
            terms.extend(_tokenize_terms(val))
    # light boost from section label words
    label = DRAFT_LABELS.get(section_key, "")
    terms.extend(_tokenize_terms(label))
    # dedupe keep order
    seen: set[str] = set()
    out: list[str] = []
    for t in terms:
        if t and t not in seen and len(t) >= 2:
            seen.add(t)
            out.append(t)
    return out[:48]


def _tokenize_terms(text: str) -> list[str]:
    parts = re.split(r"[\s,./·|/()\[\]{}:;\"'<>]+", text.lower())
    return [p for p in parts if len(p) >= 2][:24]


def _rank_citations(
    citations: list[Citation],
    *,
    terms: list[str],
    limit: int,
) -> list[Citation]:
    if not citations or limit <= 0:
        return []
    scored: list[tuple[float, Citation]] = []
    for c in citations:
        blob = f"{c.filename} {c.snippet}".lower()
        hit = 0.0
        for t in terms:
            if t in blob:
                hit += 1.0
                # filename hits count more (document-level relevance)
                if t in (c.filename or "").lower():
                    hit += 0.5
        # keep some score signal from original retrieval
        total = hit * 2.0 + float(c.score or 0.0)
        scored.append((total, c))
    scored.sort(key=lambda x: -x[0])
    # if nothing matched terms, fall back to original top scores
    if scored and scored[0][0] <= float(scored[0][1].score or 0.0) + 0.01:
        scored = sorted(
            ((float(c.score or 0.0), c) for c in citations),
            key=lambda x: -x[0],
        )
    return [c for _, c in scored[:limit]]


def _chunks_to_text(chunks: list[dict[str, str]], max_chars: int = 14000) -> str:
    parts: list[str] = []
    total = 0
    for ch in chunks:
        piece = f"[출처: {ch.get('file')} / {ch.get('location')}]\n{ch.get('text', '')}\n"
        if total + len(piece) > max_chars:
            parts.append("\n...(생략)...")
            break
        parts.append(piece)
        total += len(piece)
    return "\n".join(parts)


def _evidence_text(evidence: list[Citation], *, label: str = "근거") -> str:
    if not evidence:
        return f"({label} 없음)"
    blocks = []
    for i, c in enumerate(evidence, start=1):
        name = (c.filename or "문서").strip()
        loc = (c.location or "").strip()
        head = f"[{i}] {name}"
        if loc:
            head += f" · {loc}"
        snippet = (c.snippet or "").strip()
        blocks.append(f"{head}\n{snippet}" if snippet else head)
    return "\n\n".join(blocks)


def _format_evidence_lines(
    evidence: list[Citation] | list[dict[str, Any]],
    *,
    start: int,
) -> list[str]:
    lines: list[str] = []
    for offset, c in enumerate(evidence):
        i = start + offset
        if isinstance(c, Citation):
            name = (c.filename or "문서").strip()
            loc = (c.location or "").strip()
            line = f"[{i}] {name}"
            if loc:
                line += f" · {loc}"
            lines.append(line)
        else:
            name = str(c.get("filename") or "문서").strip()
            loc = str(c.get("location") or "").strip()
            line = f"[{i}] {name}"
            if loc:
                line += f" · {loc}"
            lines.append(line)
        lines.append("")
    return lines


def _empty_rfp() -> dict[str, Any]:
    result: dict[str, Any] = {}
    for field in RFP_FIELDS:
        if field in {
            "mandatory_requirements",
            "tech_requirements",
            "kpi",
            "evaluation_criteria",
            "submission_documents",
        }:
            result[field] = [NOT_FOUND]
        else:
            result[field] = NOT_FOUND
    return result


def _parse_rfp_json(raw: str) -> dict[str, Any]:
    cleaned = _strip_fence(raw)
    try:
        parsed = json.loads(cleaned)
    except json.JSONDecodeError:
        result = _empty_rfp()
        result["error"] = "LLM RFP JSON parse failed"
        result["raw_response"] = raw[:1000]
        return result
    result = _empty_rfp()
    for field in RFP_FIELDS:
        if field in parsed and parsed[field]:
            result[field] = parsed[field]
    result["mode"] = "llm"
    return result


def _heuristic_rfp(source: str, chunks: list[dict[str, str]]) -> dict[str, Any]:
    result = _empty_rfp()
    head = source[:2000]
    m = re.search(r"(과제명|사업명)\s*[:：]?\s*(.+)", head)
    if m:
        result["project_name"] = m.group(2).strip()[:200]
    else:
        result["project_name"] = chunks[0]["file"] if chunks else NOT_FOUND
    result["purpose"] = head[:400].replace("\n", " ")
    result["mandatory_requirements"] = [
        line.strip(" -•\t")
        for line in source.splitlines()
        if any(k in line for k in ("수행", "구축", "개발", "분석", "제안"))
    ][:8] or [NOT_FOUND]
    result["notes"] = "LLM offline — heuristic RFP parse"
    result["mode"] = "heuristic"
    return result


def _heuristic_roles(rfp: dict[str, Any], evidence: list[Citation]) -> list[dict[str, Any]]:
    ev = (
        f"{evidence[0].filename}: {evidence[0].snippet[:160]}"
        if evidence
        else NOT_FOUND
    )
    reqs = rfp.get("mandatory_requirements") or []
    related = ", ".join(reqs[:3]) if isinstance(reqs, list) else str(reqs)
    return [
        {
            "role": "Document Intelligence / Research Memory 담당",
            "reason": f"RFP 요구({related})와 문서지능·연구메모리 역량 연계 (추측 — 확인 필요)",
            "related_requirements": related or NOT_FOUND,
            "evidence": ev,
            "open_questions": "주관기관 역할 분담 확인 필요",
        }
    ]


def _heuristic_section(
    section_key: str,
    rfp: dict[str, Any],
    selected_role: dict[str, Any],
    research: list[Citation],
    reference: list[Citation],
) -> str:
    label = DRAFT_LABELS.get(section_key, section_key)
    if section_key == "compliance_notes":
        if reference:
            return (
                f"- 연구시설·장비 현물 계상 및 내부공간 임차료 계상 기준 확인 필요 "
                f"({reference[0].filename})"
            )
        return "- 운영요령·참고규정 근거 없음 — 확인 필요"
    if section_key == "necessity":
        return (
            f"- {rfp.get('project_name', NOT_FOUND)} 참여 필요성 구체화 필요\n"
            f"- RFP 목적과 기존 연구 성과 연계 서술 필요"
        )
    if section_key == "center_role":
        role = selected_role.get("role", NOT_FOUND)
        return f"- 담당 역할: {role}\n- 역할 범위·경계 상세화 필요"
    if section_key == "open_questions":
        return "- LLM 연결 후 섹션별 초안 재생성 필요"
    if research:
        return f"({research[0].filename}) {research[0].snippet[:240]}"
    if reference:
        return f"({reference[0].filename}) {reference[0].snippet[:240]}"
    return f"확인 필요 — {label} 섹션 근거 없음"


def _parse_role_list(raw: str) -> list[dict[str, Any]]:
    cleaned = _strip_fence(raw)
    try:
        parsed = json.loads(cleaned)
        if isinstance(parsed, list):
            return parsed[:3]
    except json.JSONDecodeError:
        pass
    return [
        {
            "role": "확인 필요 (응답 해석 실패)",
            "reason": "LLM 응답을 JSON으로 해석하지 못했습니다.",
            "related_requirements": NOT_FOUND,
            "evidence": NOT_FOUND,
            "open_questions": raw[:500],
        }
    ]


def _parse_section_text(raw: str, section_key: str) -> str:
    cleaned = _strip_fence(raw)
    try:
        parsed = json.loads(cleaned)
        if isinstance(parsed, dict):
            val = parsed.get(section_key)
            if val is None and len(parsed) == 1:
                val = next(iter(parsed.values()))
            if isinstance(val, str) and val.strip():
                return val.strip()
            if val is not None:
                return str(val).strip()
    except json.JSONDecodeError:
        pass
    # model sometimes returns bare prose
    text = cleaned.strip()
    return text if text else NOT_FOUND


def _strip_fence(text: str) -> str:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.strip("`")
        if cleaned.lower().startswith("json"):
            cleaned = cleaned[4:]
        cleaned = cleaned.strip()
    return cleaned
