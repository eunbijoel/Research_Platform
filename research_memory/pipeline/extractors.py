from __future__ import annotations

import re
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import pandas as pd
from docx import Document
from pypdf import PdfReader

from research_memory.schema import TextChunk


def extract_chunks(path: Path) -> tuple[str, list[TextChunk], str]:
    """
    Return (file_type, chunks, error).
    error is empty on success; chunks may still be empty for blank docs.
    """
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            return "pdf", _pdf(path), ""
        if ext == ".docx":
            return "docx", _docx(path), ""
        if ext in {".txt", ".md"}:
            return ext.lstrip("."), _plain(path), ""
        if ext == ".csv":
            return "csv", _csv(path), ""
        if ext in {".xlsx", ".xls"}:
            return "excel", _excel(path), ""
        if ext == ".hwpx":
            return "hwpx", _hwpx(path), ""
        if ext == ".hwp":
            return "hwp", [], (
                "구형 .hwp(한글 바이너리)는 아직 지원하지 않습니다. "
                "한글에서 .hwpx로 다시 저장하거나 PDF/DOCX로 변환해 주세요."
            )
        return "unknown", [], f"Unsupported extension: {ext}"
    except Exception as exc:  # noqa: BLE001 — surface parse failures to ingest
        return ext.lstrip(".") or "unknown", [], str(exc)


def extract_chunks_from_bytes(data: bytes, filename: str) -> tuple[str, list[TextChunk], str]:
    """Extract chunks from uploaded file bytes using extension inferred from filename."""
    suffix = Path(filename).suffix or ".bin"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(data)
        temp_path = Path(tmp.name)
    try:
        return extract_chunks(temp_path)
    finally:
        temp_path.unlink(missing_ok=True)


def _pdf(path: Path) -> list[TextChunk]:
    reader = PdfReader(str(path))
    chunks: list[TextChunk] = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            chunks.append(TextChunk(text=text, location=f"{i}페이지", page=i))
    return chunks


def _docx(path: Path) -> list[TextChunk]:
    doc = Document(str(path))
    chunks: list[TextChunk] = []
    para_n = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        para_n += 1
        chunks.append(TextChunk(text=text, location=f"문단 {para_n}"))
    for t_i, table in enumerate(doc.tables, start=1):
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                chunks.append(TextChunk(text=row_text, location=f"표 {t_i}"))
    return chunks


def _plain(path: Path) -> list[TextChunk]:
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    if not text:
        return []
    return [TextChunk(text=text, location="전체")]


def _csv(path: Path) -> list[TextChunk]:
    df = pd.read_csv(path)
    preview = df.head(40).to_csv(index=False)
    summary = f"CSV columns: {', '.join(map(str, df.columns))} | rows={len(df)}\n{preview}"
    return [TextChunk(text=summary, location="시트 preview")]


def _excel(path: Path) -> list[TextChunk]:
    book = pd.read_excel(path, sheet_name=None)
    chunks: list[TextChunk] = []
    for name, df in book.items():
        preview = df.head(30).to_csv(index=False)
        text = f"Sheet={name} columns={list(df.columns)} rows={len(df)}\n{preview}"
        chunks.append(TextChunk(text=text, location=f"시트 {name}"))
    return chunks


def _hwpx(path: Path) -> list[TextChunk]:
    """Best-effort HWPX (zip+xml) text extraction without hwpilot."""
    head = path.read_bytes()[:8]
    # Classic .hwp is OLE compound; real .hwpx is a ZIP package.
    if head[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
        raise ValueError(
            "확장자는 .hwpx지만 실제 파일은 구형 .hwp(한글 바이너리)입니다. "
            "한글에서 [다른 이름으로 저장] → .hwpx 또는 PDF/DOCX로 변환 후 다시 올려 주세요."
        )
    if head[:2] != b"PK":
        raise ValueError(
            "유효한 .hwpx(ZIP) 파일이 아닙니다. "
            "파일이 손상되었거나 확장자만 .hwpx인 다른 형식일 수 있습니다."
        )
    texts: list[str] = []
    with zipfile.ZipFile(path) as zf:
        section_names = sorted(
            n for n in zf.namelist() if re.search(r"Contents/section\d+\.xml$", n)
        )
        if not section_names:
            section_names = [n for n in zf.namelist() if n.endswith(".xml")]
        for name in section_names:
            raw = zf.read(name)
            root = ET.fromstring(raw)
            parts = [
                (node.text or "").strip()
                for node in root.iter()
                if node.text and node.text.strip()
            ]
            joined = "\n".join(parts).strip()
            if joined:
                texts.append(joined)
    chunks = [
        TextChunk(text=t, location=f"section {i}")
        for i, t in enumerate(texts, start=1)
    ]
    return chunks
